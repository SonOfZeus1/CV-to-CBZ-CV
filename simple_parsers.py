import logging
import re
import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path: str):
    """
    Extracts text from a PDF file using PyMuPDF.
    Returns a tuple (text, ocr_applied).
    ocr_applied is always False here as we are doing simple extraction.
    """
    text = ""
    try:
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
    except Exception as e:
        logger.error(f"Error reading PDF {pdf_path}: {e}")
        return "", False
        
    return text, False

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts text from a DOCX file.
    """
    text = []
    try:
        doc = Document(docx_path)
        
        # 1. Paragraphs
        for para in doc.paragraphs:
            text.append(para.text)
            
        # 2. Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text.append(para.text)
                        
        # 3. Headers and Footers
        for section in doc.sections:
            # Header
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        text.append(para.text)
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    text.append(para.text)
            # Footer
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        text.append(para.text)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    text.append(para.text)
                                    
    except Exception as e:
        logger.error(f"Error reading DOCX {docx_path}: {e}")
        return ""
        
    return "\n".join(text)

def heuristic_parse_contact(text: str) -> dict:
    """
    Extracts contact info (phone) using regex.
    Returns a dict with 'phone'.
    """
    contact_info = {"phone": ""}
    if not text:
        return contact_info
        
    # Phone Regex (North American formats mostly)
    # Matches: (123) 456-7890, 123-456-7890, 123.456.7890, +1 123 456 7890
    phone_pattern = r'(?:(?:\+?1\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?'
    
    # Simpler regex to catch more cases, then clean up
    # Look for groups of digits that look like a phone number
    # e.g. 514 123 4567
    simple_phone_pattern = r'[\+\(]?[0-9][0-9 .\-\(\)]{8,}[0-9]'
    
    matches = re.findall(simple_phone_pattern, text)
    
    best_phone = ""
    for match in matches:
        # Clean up
        digits = re.sub(r'\D', '', match)
        if 10 <= len(digits) <= 15:
            # It's likely a phone number
            best_phone = match.strip()
            break
            
    contact_info["phone"] = best_phone
    return contact_info
